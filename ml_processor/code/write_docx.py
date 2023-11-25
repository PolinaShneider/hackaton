from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import timedelta

def convert_to_hh_mm_ss(time_float):
    # Convert time in float format to hh:mm:ss string
    total_seconds = int(time_float)
    hours, remainder = divmod(total_seconds, 3600)
    minutes, seconds = divmod(remainder, 60)

    return f"{hours:02}:{minutes:02}:{seconds:02}"

# не добавили тамйкоды для саммари
# отступы странные
def create_docx_from_list(data):
    doc = Document()

    title = doc.add_heading("Конспект лекции\n", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    bullet_list = doc.add_paragraph()
    term_counter = 1

    for item in data:
        if len(item) == 3:
            # Add a string as regular text with an interval
            summary, summary_start, summary_end = item
            
            timestart_str = convert_to_hh_mm_ss(summary_start)
            timeend_str = convert_to_hh_mm_ss(summary_end)
            
            if summary.endswith('.'):
                summary = summary[:-1]
       
            summary = summary.split('\n')[0].capitalize()

            bullet_list.add_run(f"{summary} ({timestart_str} - {timeend_str})\n\n")
        else:
            # skip filename
            term, definition, timestart, timeend = item
            term = term.capitalize()
            definition = definition.capitalize()
            timestart_str = convert_to_hh_mm_ss(timestart)
            timeend_str = convert_to_hh_mm_ss(timeend)
            
            # Add an element to the bullet list with formatted content
            bullet_list.add_run(f"{term_counter}. {term} ").bold = True
            bullet_list.add_run(f"({timestart_str} - {timeend_str})\n{definition}").bold = False

            term_counter += 1

            bullet_list.add_run("\n\n")  # Add a line break

    return doc
